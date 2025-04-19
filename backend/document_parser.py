from abc import ABC, abstractmethod
import fitz  # PyMuPDF
import pandas as pd
import docx
import markdown
import os
import logging

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

class DocumentParser(ABC):
    """文档解析的抽象基类，定义统一接口"""
    
    @abstractmethod
    def parse(self, file_path):
        """
        解析文档，提取文本内容和元数据
        
        Args:
            file_path: 文档路径
            
        Returns:
            dict: 包含解析后的文本、元数据及提取的实体
        """
        pass
    
    def extract_entities(self, text):
        """
        从文本中提取实体，后续将替换为NER模型
        
        Args:
            text: 待提取实体的文本
            
        Returns:
            list: 提取的实体列表
        """
        # 临时实现，后续会被实际的NER模型替换
        return []


class PDFParser(DocumentParser):
    """PDF文档解析器"""
    
    def parse(self, file_path):
        """解析PDF文档"""
        try:
            doc = fitz.open(file_path)
            text = ""
            # 提取文本内容
            for page in doc:
                text += page.get_text()
            
            # 提取元数据
            metadata = {
                "file_type": "pdf",
                "page_count": len(doc),
                "title": doc.metadata.get("title", ""),
                "author": doc.metadata.get("author", ""),
                "creation_date": doc.metadata.get("creationDate", "")
            }
            
            # 清洗文本
            cleaned_text = self._clean_text(text)
            
            # 提取实体
            entities = self.extract_entities(cleaned_text)
            
            return {
                "text": cleaned_text, 
                "meta": metadata, 
                "entities": entities,
                "source_file": os.path.basename(file_path)
            }
        except Exception as e:
            logger.error(f"解析PDF文件 {file_path} 时出错: {str(e)}")
            return {"error": str(e)}
    
    def _clean_text(self, text):
        """清洗文本，去除特殊字符、标准化术语等"""
        # 去除重复空行
        cleaned = "\n".join([line for line in text.splitlines() if line.strip()])
        
        # 标准化术语
        term_mapping = {
            "持仓": "持仓量",
            "保证金": "保证金比例",
            # 可添加更多术语标准化映射
        }
        
        for old_term, new_term in term_mapping.items():
            cleaned = cleaned.replace(old_term, new_term)
            
        return cleaned


class WordParser(DocumentParser):
    """Word文档解析器"""
    
    def parse(self, file_path):
        """解析Word文档"""
        try:
            doc = docx.Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            
            # 提取元数据
            metadata = {
                "file_type": "docx",
                "paragraph_count": len(doc.paragraphs),
                "source_file": os.path.basename(file_path)
            }
            
            # 提取表格内容
            tables_content = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables_content.append(table_data)
            
            # 清洗文本
            cleaned_text = self._clean_text(text)
            
            # 提取实体
            entities = self.extract_entities(cleaned_text)
            
            return {
                "text": cleaned_text,
                "tables": tables_content,
                "meta": metadata,
                "entities": entities,
                "source_file": os.path.basename(file_path)
            }
        except Exception as e:
            logger.error(f"解析Word文件 {file_path} 时出错: {str(e)}")
            return {"error": str(e)}
    
    def _clean_text(self, text):
        """清洗文本"""
        # 类似PDF清洗逻辑
        cleaned = "\n".join([line for line in text.splitlines() if line.strip()])
        
        # 标准化术语
        term_mapping = {
            "持仓": "持仓量",
            "保证金": "保证金比例",
        }
        
        for old_term, new_term in term_mapping.items():
            cleaned = cleaned.replace(old_term, new_term)
            
        return cleaned


class ExcelParser(DocumentParser):
    """Excel文档解析器"""
    
    def parse(self, file_path):
        """解析Excel文档"""
        try:
            # 使用openpyxl引擎读取Excel
            excel_data = pd.read_excel(file_path, engine="openpyxl", sheet_name=None)
            
            # 将所有sheet数据转为字典
            sheet_data = {}
            for sheet_name, df in excel_data.items():
                # 尝试将DataFrame转为记录列表
                sheet_data[sheet_name] = df.to_dict("records")
            
            metadata = {
                "file_type": "excel",
                "sheet_count": len(excel_data),
                "sheet_names": list(excel_data.keys()),
                "source_file": os.path.basename(file_path)
            }
            
            # 提取实体（从所有sheet中）
            all_entities = []
            for sheet_name, records in sheet_data.items():
                sheet_text = str(records)  # 简单转为文本
                sheet_entities = self.extract_entities(sheet_text)
                # 为每个实体添加来源sheet信息
                for entity in sheet_entities:
                    entity["sheet"] = sheet_name
                all_entities.extend(sheet_entities)
            
            return {
                "data": sheet_data,
                "meta": metadata,
                "entities": all_entities,
                "source_file": os.path.basename(file_path)
            }
        except Exception as e:
            logger.error(f"解析Excel文件 {file_path} 时出错: {str(e)}")
            return {"error": str(e)}


class MarkdownParser(DocumentParser):
    """Markdown文档解析器"""
    
    def parse(self, file_path):
        """解析Markdown文档"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                md_content = f.read()
            
            # 将Markdown转为HTML，再提取纯文本
            html = markdown.markdown(md_content)
            # 简单地移除HTML标签（实际中可能需要更复杂的HTML解析）
            text = html.replace("<p>", "\n").replace("</p>", "")
            text = text.replace("<h1>", "\n# ").replace("</h1>", "")
            text = text.replace("<h2>", "\n## ").replace("</h2>", "")
            # ... 更多标签处理
            
            metadata = {
                "file_type": "markdown",
                "source_file": os.path.basename(file_path)
            }
            
            # 清洗文本
            cleaned_text = self._clean_text(text)
            
            # 提取实体
            entities = self.extract_entities(cleaned_text)
            
            return {
                "text": cleaned_text,
                "meta": metadata,
                "entities": entities,
                "source_file": os.path.basename(file_path)
            }
        except Exception as e:
            logger.error(f"解析Markdown文件 {file_path} 时出错: {str(e)}")
            return {"error": str(e)}
    
    def _clean_text(self, text):
        """清洗文本"""
        # 去除多余的空行
        cleaned = "\n".join([line for line in text.splitlines() if line.strip()])
        
        # 标准化术语
        term_mapping = {
            "持仓": "持仓量",
            "保证金": "保证金比例",
        }
        
        for old_term, new_term in term_mapping.items():
            cleaned = cleaned.replace(old_term, new_term)
            
        return cleaned


class ParserFactory:
    """文档解析器工厂类，根据文件类型创建相应的解析器"""
    
    @staticmethod
    def create_parser(file_path):
        """
        根据文件路径创建对应的解析器
        
        Args:
            file_path: 文件路径
            
        Returns:
            DocumentParser: 对应类型的解析器实例
        """
        extension = os.path.splitext(file_path)[1].lower()
        
        if extension == ".pdf":
            return PDFParser()
        elif extension in [".docx", ".doc"]:
            return WordParser()
        elif extension in [".xlsx", ".xls"]:
            return ExcelParser()
        elif extension in [".md", ".markdown"]:
            return MarkdownParser()
        else:
            logger.warning(f"不支持的文件类型: {extension}")
            raise ValueError(f"不支持的文件类型: {extension}")


# 使用示例
if __name__ == "__main__":
    # 假设有测试文件
    test_files = [
        "samples/contract.pdf",
        "samples/rules.docx",
        "samples/margins.xlsx",
        "samples/guide.md"
    ]
    
    for file_path in test_files:
        try:
            parser = ParserFactory.create_parser(file_path)
            result = parser.parse(file_path)
            print(f"解析 {file_path} 成功:")
            print(f"元数据: {result['meta']}")
            print(f"实体数量: {len(result['entities'])}")
            # 如果是文本型文档，打印部分文本
            if 'text' in result:
                print(f"文本预览: {result['text'][:200]}...")
            print("-" * 50)
        except Exception as e:
            print(f"解析 {file_path} 失败: {str(e)}")